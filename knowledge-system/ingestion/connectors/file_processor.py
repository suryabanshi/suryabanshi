"""
Local file system connector.
Handles: PDF, Word (.docx), Excel (.xlsx/.csv), Markdown, plain text, images.
"""
from __future__ import annotations

import hashlib
import mimetypes
import os
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterator, Optional

# Optional imports — skip gracefully if library not installed
try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False


SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv",
    ".md", ".txt", ".rst", ".json", ".yaml", ".yml",
    ".html", ".htm",
}


@dataclass
class RawDocument:
    title: str
    raw_content: str
    content_hash: str
    mime_type: str
    file_size_bytes: int
    source_url: str                # absolute path or URL
    source_metadata: dict
    created_at_source: Optional[datetime] = None
    external_id: Optional[str] = None


def sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


class LocalFileConnector:
    """
    Recursively walks a directory and yields RawDocument objects.

    Usage:
        connector = LocalFileConnector("/path/to/vault")
        for doc in connector.extract():
            ingest(doc)
    """

    def __init__(self, root_path: str, recursive: bool = True):
        self.root = Path(root_path)
        self.recursive = recursive

    def extract(self) -> Iterator[RawDocument]:
        pattern = "**/*" if self.recursive else "*"
        for path in self.root.glob(pattern):
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
                doc = self._process_file(path)
                if doc:
                    yield doc

    def _process_file(self, path: Path) -> Optional[RawDocument]:
        ext = path.suffix.lower()
        mime, _ = mimetypes.guess_type(str(path))
        size = path.stat().st_size
        mtime = datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc)

        content = self._read_content(path, ext)
        if not content or not content.strip():
            return None

        return RawDocument(
            title=path.stem,
            raw_content=content,
            content_hash=sha256(content),
            mime_type=mime or "text/plain",
            file_size_bytes=size,
            source_url=str(path.resolve()),
            source_metadata={
                "filename": path.name,
                "extension": ext,
                "directory": str(path.parent),
                "modified_at": mtime.isoformat(),
            },
            created_at_source=mtime,
            external_id=str(path.resolve()),
        )

    def _read_content(self, path: Path, ext: str) -> str:
        try:
            if ext == ".pdf":
                return self._read_pdf(path)
            if ext in (".docx", ".doc"):
                return self._read_docx(path)
            if ext in (".xlsx", ".xls"):
                return self._read_xlsx(path)
            if ext == ".csv":
                return self._read_csv(path)
            # Text-based formats
            return path.read_text(encoding="utf-8", errors="replace")
        except Exception as exc:  # noqa: BLE001
            print(f"[FileConnector] Skipping {path}: {exc}")
            return ""

    @staticmethod
    def _read_pdf(path: Path) -> str:
        if not HAS_PDF:
            raise ImportError("pdfplumber not installed — run: pip install pdfplumber")
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n\n".join(pages)

    @staticmethod
    def _read_docx(path: Path) -> str:
        if not HAS_DOCX:
            raise ImportError("python-docx not installed — run: pip install python-docx")
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    @staticmethod
    def _read_xlsx(path: Path) -> str:
        if HAS_PANDAS:
            sheets = []
            xls = pd.ExcelFile(str(path))
            for sheet in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet)
                sheets.append(f"[Sheet: {sheet}]\n{df.to_string(index=False)}")
            return "\n\n".join(sheets)
        if not HAS_XLSX:
            raise ImportError("openpyxl not installed — run: pip install openpyxl")
        wb = openpyxl.load_workbook(str(path), data_only=True)
        rows = []
        for sheet in wb.worksheets:
            rows.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                rows.append("\t".join(str(c) if c is not None else "" for c in row))
        return "\n".join(rows)

    @staticmethod
    def _read_csv(path: Path) -> str:
        if HAS_PANDAS:
            df = pd.read_csv(str(path))
            return df.to_string(index=False)
        return path.read_text(encoding="utf-8", errors="replace")
